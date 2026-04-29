"""DOCX source adapter (Gap 3 slice 2).

Turns a Word ``.docx`` file on disk (or in-memory bytes) into an
``IngestionEnvelope``: one chunk per heading-delimited section,
falling back to one chunk per paragraph group when no headings are
present. Each chunk carries a locator like ``section=2,paragraphs=5-12``
so operators can trace which chunk of the document produced which
turn.

Design rules (mirror ``pdf.py``):

* **Lazy dep import.** ``python-docx`` is declared as
  ``lifeform-ingestion[docx]`` optional; unused installs pay no
  cost. ``ImportError`` surfaces a clear message naming the extra.
* **Tables + inline images ignored.** slice 2 extracts paragraph
  text only. Tables and floating images become a structured
  ``partial_failure`` (``section_has_unextracted_table``) so the
  operator sees what was skipped. slice 2b can add table support.
* **Structural chunking.** Sections are delimited by paragraphs
  whose ``style.name`` starts with "Heading" (case-insensitive).
  When no headings exist the whole document is a single section,
  sub-chunked by ``max_chunk_chars`` so long docs do not overflow
  the turn budget.
* **Bounded paragraph count.** ``max_paragraphs`` defaults to
  2000 to catch runaway inputs.
* **No network.** Only file paths and bytes; no URL fetch.
"""

from __future__ import annotations

import hashlib
import io
import pathlib
import time
import zipfile
from typing import Any

from lifeform_ingestion.envelope import (
    IngestionChunk,
    IngestionComplianceProfile,
    IngestionEnvelope,
    IngestionProvenance,
    IngestionSourceKind,
)


DEFAULT_MAX_PARAGRAPHS: int = 2000
DEFAULT_MAX_CHUNK_CHARS: int = 4096


class DocxIngestionError(ValueError):
    """Raised when a DOCX cannot be opened / parsed / exceeds caps."""


def _load_docx_module() -> Any:
    """Lazy-import ``docx`` (python-docx) with a clear error message."""
    try:
        import docx  # type: ignore
    except ImportError as exc:  # pragma: no cover - depends on install
        raise DocxIngestionError(
            "DOCX ingestion requires the python-docx dependency. Install "
            "with 'pip install lifeform-ingestion[docx]' or add "
            "python-docx>=1.0 to your environment."
        ) from exc
    return docx


def _paragraph_is_heading(paragraph: Any) -> bool:
    """Return True iff the paragraph's style name begins with 'Heading'.

    python-docx attaches ``paragraph.style`` which may be None for
    default-styled paragraphs. ``getattr`` chains defensively so a
    stub paragraph without a style doesn't crash.
    """
    style = getattr(paragraph, "style", None)
    if style is None:
        return False
    name = getattr(style, "name", None)
    if not name:
        return False
    return str(name).lower().startswith("heading")


def _extract_sections(
    data: bytes, *, max_paragraphs: int
) -> tuple[tuple[str, tuple[int, int], bool], ...]:
    """Extract ``(section_text, (first_para, last_para), had_table)`` tuples.

    Raises ``DocxIngestionError`` on document-level failure.
    Per-section errors do not abort; they become a bool ``had_table``
    tag so the caller can mark partial failure.

    ``first_para`` / ``last_para`` are 1-based inclusive indices
    into the document's paragraph stream, useful for locator
    strings.
    """
    docx_mod = _load_docx_module()
    try:
        document = docx_mod.Document(io.BytesIO(data))
    except Exception as exc:  # noqa: BLE001 - re-raised typed
        raise DocxIngestionError(
            f"python-docx failed to open DOCX: {type(exc).__name__}: {exc}"
        ) from exc
    paragraphs = list(document.paragraphs)
    if len(paragraphs) > max_paragraphs:
        raise DocxIngestionError(
            f"DOCX has {len(paragraphs)} paragraphs, exceeds "
            f"max_paragraphs={max_paragraphs}. Split the document or "
            f"raise the limit."
        )
    # Observation about python-docx: ``document.paragraphs`` includes
    # only top-level paragraphs; text inside tables is surfaced via
    # ``document.tables[*].rows[*].cells[*].paragraphs``. Slice 2
    # deliberately does not deep-traverse tables \u2014 we flag their
    # presence with had_table=True and leave structured extraction
    # to slice 2b.
    tables_present = bool(getattr(document, "tables", None))
    sections: list[tuple[str, tuple[int, int], bool]] = []
    current_text: list[str] = []
    current_first = 1
    index_1_based = 0
    # had_table applies to a section if there's any table between
    # its first and last paragraph \u2014 but python-docx does not
    # expose paragraph \u2194 table order cheaply. Slice 2 marks
    # every section with ``had_table`` equal to the document-level
    # tables_present flag; this is coarse but honest (the operator
    # sees "this document has tables we didn't extract") and
    # doesn't pretend to have per-section fidelity.
    for paragraph in paragraphs:
        index_1_based += 1
        text = (paragraph.text or "").strip()
        if _paragraph_is_heading(paragraph):
            # Close current section (if any).
            if current_text:
                sections.append(
                    (
                        "\n\n".join(current_text),
                        (current_first, index_1_based - 1),
                        tables_present,
                    )
                )
            # Start new section with the heading text as the first
            # paragraph of the new section.
            current_text = [text] if text else []
            current_first = index_1_based
            continue
        if text:
            current_text.append(text)
    # Flush the last section.
    if current_text:
        sections.append(
            (
                "\n\n".join(current_text),
                (current_first, index_1_based),
                tables_present,
            )
        )
    return tuple(sections)


def envelope_from_docx_bytes(
    data: bytes,
    *,
    source_uri: str,
    uploader: str = "system",
    upload_ts_ms: int | None = None,
    envelope_id: str | None = None,
    compliance_profile: IngestionComplianceProfile = IngestionComplianceProfile.FORCED,
    max_paragraphs: int = DEFAULT_MAX_PARAGRAPHS,
    max_chunk_chars: int = DEFAULT_MAX_CHUNK_CHARS,
) -> IngestionEnvelope:
    """Build an envelope from raw DOCX bytes.

    DOCX is a zip container; we do a fast magic-bytes check
    (``PK\\x03\\x04``) before handing it to python-docx so a
    mis-named text / PDF file fails loud instead of producing
    a confusing exception trace.
    """
    if not data:
        raise DocxIngestionError(
            f"envelope_from_docx_bytes: input for {source_uri!r} is empty"
        )
    if not data.startswith(b"PK\x03\x04"):
        raise DocxIngestionError(
            f"envelope_from_docx_bytes: {source_uri!r} does not start with "
            f"ZIP magic bytes (PK\\x03\\x04); DOCX is a ZIP container."
        )
    # Confirm the zip contains ``word/document.xml`` \u2014 a renamed
    # but otherwise valid ZIP should not silently accept.
    try:
        with zipfile.ZipFile(io.BytesIO(data)) as zf:
            if "word/document.xml" not in zf.namelist():
                raise DocxIngestionError(
                    f"envelope_from_docx_bytes: {source_uri!r} is a ZIP but "
                    f"does not contain 'word/document.xml'; not a DOCX file."
                )
    except zipfile.BadZipFile as exc:
        raise DocxIngestionError(
            f"envelope_from_docx_bytes: {source_uri!r} is not a valid ZIP: {exc}"
        ) from exc
    if max_paragraphs <= 0:
        raise DocxIngestionError(
            f"envelope_from_docx_bytes: max_paragraphs must be > 0, "
            f"got {max_paragraphs!r}"
        )
    sections = _extract_sections(data, max_paragraphs=max_paragraphs)
    if not sections:
        raise DocxIngestionError(
            f"envelope_from_docx_bytes: {source_uri!r} produced no extractable "
            f"paragraphs; nothing to ingest."
        )
    integrity_hash = hashlib.sha256(data).hexdigest()
    if envelope_id is None:
        envelope_id = f"docx:{integrity_hash[:12]}"
    if upload_ts_ms is None:
        upload_ts_ms = int(time.time() * 1000.0)
    chunks: list[IngestionChunk] = []
    failures: list[str] = []
    emitted_any_text = False
    for section_index, (section_text, (first_para, last_para), had_table) in enumerate(sections):
        base_chunk_id = f"{envelope_id}:section:{section_index:04d}"
        base_locator = (
            f"section={section_index + 1}/{len(sections)},"
            f"paragraphs={first_para}-{last_para}"
        )
        section_parse_error = (
            "section_has_unextracted_table"
            if had_table
            else ""
        )
        # Sub-chunk a long section.
        cursor = 0
        sub_index = 0
        section_emitted = False
        while cursor < len(section_text):
            end = min(cursor + max_chunk_chars, len(section_text))
            segment = section_text[cursor:end]
            if segment.strip():
                sub_chunk_id = (
                    base_chunk_id
                    if sub_index == 0
                    else f"{base_chunk_id}:sub:{sub_index:02d}"
                )
                sub_locator = (
                    base_locator
                    if cursor == 0 and end == len(section_text)
                    else f"{base_locator},offset={cursor}-{end}"
                )
                # Only the FIRST sub-chunk carries the "table skipped"
                # parse_error so the failure shows up exactly once per
                # section. Remaining sub-chunks are clean successes.
                sub_parse_error = section_parse_error if sub_index == 0 else ""
                chunks.append(
                    IngestionChunk(
                        chunk_id=sub_chunk_id,
                        text=segment,
                        locator=sub_locator,
                        confidence=0.85 if sub_parse_error else 1.0,
                        parse_error=sub_parse_error,
                    )
                )
                if sub_parse_error:
                    failures.append(sub_chunk_id)
                section_emitted = True
                emitted_any_text = True
                sub_index += 1
            cursor = end
        if not section_emitted:
            # Empty section (only headings, no body) \u2014 record as a
            # structured partial failure, same discipline as PDF's
            # page_empty.
            chunks.append(
                IngestionChunk(
                    chunk_id=base_chunk_id,
                    text="",
                    locator=base_locator,
                    confidence=0.0,
                    parse_error="section_empty",
                )
            )
            failures.append(base_chunk_id)
    if not emitted_any_text:
        raise DocxIngestionError(
            f"envelope_from_docx_bytes: {source_uri!r} produced {len(sections)} "
            f"sections but none had extractable text."
        )
    provenance = IngestionProvenance(
        uploader=uploader,
        upload_ts_ms=upload_ts_ms,
        source_uri=source_uri,
        integrity_hash=integrity_hash,
    )
    return IngestionEnvelope(
        envelope_id=envelope_id,
        # DOCX maps to BOOK semantically (authored document) rather
        # than CORPUS. Downstream reflection writeback can key off
        # ``source_kind=BOOK`` + ``source_uri`` to attribute
        # ingested records back to specific documents.
        source_kind=IngestionSourceKind.BOOK,
        chunks=tuple(chunks),
        provenance=provenance,
        compliance_profile=compliance_profile,
        partial_failures=tuple(failures),
    )


def envelope_from_docx_file(
    path: str | pathlib.Path,
    *,
    uploader: str = "system",
    upload_ts_ms: int | None = None,
    envelope_id: str | None = None,
    compliance_profile: IngestionComplianceProfile = IngestionComplianceProfile.FORCED,
    max_paragraphs: int = DEFAULT_MAX_PARAGRAPHS,
    max_chunk_chars: int = DEFAULT_MAX_CHUNK_CHARS,
) -> IngestionEnvelope:
    """Build an envelope from a ``.docx`` file on disk."""
    p = pathlib.Path(path)
    if not p.is_file():
        raise DocxIngestionError(
            f"envelope_from_docx_file: {p!s} is not a regular file"
        )
    data = p.read_bytes()
    return envelope_from_docx_bytes(
        data,
        source_uri=f"file://{p.resolve()}",
        uploader=uploader,
        upload_ts_ms=upload_ts_ms,
        envelope_id=envelope_id,
        compliance_profile=compliance_profile,
        max_paragraphs=max_paragraphs,
        max_chunk_chars=max_chunk_chars,
    )


__all__ = [
    "DEFAULT_MAX_CHUNK_CHARS",
    "DEFAULT_MAX_PARAGRAPHS",
    "DocxIngestionError",
    "envelope_from_docx_bytes",
    "envelope_from_docx_file",
]
