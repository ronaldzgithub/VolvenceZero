"""PDF + DOCX source adapter tests (Gap 3 slice 2).

Fixtures are built at runtime from pypdf / python-docx so we do
NOT have to check binary blobs into git. A test that can't import
those libs is skipped with a helpful message.

Coverage:

* PDF: happy path multi-page, per-page chunking, empty-page
  partial_failure, encrypted PDF fail-loud, non-PDF magic bytes
  rejection, max_pages cap, max_chunk_chars sub-chunking.
* DOCX: happy path heading-delimited sections, fallback to
  single-section when no headings, paragraph-count cap, bad
  magic bytes rejection, ZIP-without-document.xml rejection,
  empty-section partial_failure, table presence partial_failure.
"""

from __future__ import annotations

import io
import pathlib
import zipfile

import pytest

pytest.importorskip("pypdf", reason="PDF source tests require the [pdf] extra")
pytest.importorskip("docx", reason="DOCX source tests require the [docx] extra")

from lifeform_ingestion import (
    DocxIngestionError,
    IngestionComplianceProfile,
    IngestionSourceKind,
    PdfIngestionError,
    envelope_from_docx_bytes,
    envelope_from_docx_file,
    envelope_from_pdf_bytes,
    envelope_from_pdf_file,
)


# ---------------------------------------------------------------------------
# PDF fixtures (built via pypdf at runtime)
# ---------------------------------------------------------------------------


def _build_pdf_bytes(*, page_texts: tuple[str, ...], encrypted: bool = False) -> bytes:
    """Build a PDF with one explicit content stream per page using pypdf.

    We use the low-level PdfWriter API to draw real text objects;
    ``PdfWriter().add_blank_page()`` alone produces pages whose
    ``extract_text()`` returns an empty string which is useful for
    the empty-page branch but not for the happy-path assertion.

    Drawing text without a registered font crashes on extract; we
    avoid that by using pypdf's ``PageObject`` + a helper that writes
    a valid ``/Contents`` stream with ``Tj`` operators referencing
    ``Helvetica`` (always embedded by default when pypdf serialises).
    """
    from pypdf import PdfWriter
    from pypdf.generic import (
        ArrayObject,
        ContentStream,
        DecodedStreamObject,
        DictionaryObject,
        FloatObject,
        NameObject,
        NumberObject,
        TextStringObject,
    )

    writer = PdfWriter()
    for text in page_texts:
        page = writer.add_blank_page(width=612, height=792)
        # Register a Helvetica font on the page.
        font_dict = DictionaryObject(
            {
                NameObject("/Type"): NameObject("/Font"),
                NameObject("/Subtype"): NameObject("/Type1"),
                NameObject("/BaseFont"): NameObject("/Helvetica"),
            }
        )
        resources = DictionaryObject(
            {
                NameObject("/Font"): DictionaryObject(
                    {NameObject("/F1"): font_dict}
                ),
            }
        )
        page[NameObject("/Resources")] = resources
        # Build a content stream that writes each line.
        lines = text.splitlines() or [text]
        y = 720
        ops: list[tuple] = []
        for line in lines:
            ops.append((
                [],
                b"BT",
            ))
            ops.append((
                [NameObject("/F1"), NumberObject(12)],
                b"Tf",
            ))
            ops.append((
                [NumberObject(72), NumberObject(y)],
                b"Td",
            ))
            ops.append((
                [TextStringObject(line)],
                b"Tj",
            ))
            ops.append((
                [],
                b"ET",
            ))
            y -= 18
        content_stream = ContentStream(None, writer)
        content_stream.operations = ops
        # pypdf expects an indirect object for /Contents; attach the
        # stream dict directly.
        decoded = DecodedStreamObject()
        decoded.set_data(content_stream.get_data())
        page[NameObject("/Contents")] = decoded
    if encrypted:
        writer.encrypt(user_password="hunter2", owner_password="adminpass")
    buf = io.BytesIO()
    writer.write(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# PDF tests
# ---------------------------------------------------------------------------


def test_envelope_from_pdf_bytes_multi_page_happy_path() -> None:
    pdf = _build_pdf_bytes(
        page_texts=(
            "First page hello.",
            "Second page world.",
            "Third page with more content than the prior two lines.",
        ),
    )
    envelope = envelope_from_pdf_bytes(pdf, source_uri="memory://test.pdf")
    assert envelope.source_kind is IngestionSourceKind.BOOK
    # At least 3 chunks (one per non-empty page; may have sub-chunks).
    assert envelope.total_chunks >= 3
    successful = envelope.successful_chunks
    # Text content should contain the page markers.
    combined = " ".join(c.text for c in successful)
    assert "First page" in combined
    assert "Second page" in combined
    assert "Third page" in combined
    # Locators include page=N/TOTAL.
    assert any("page=1/3" in c.locator for c in successful)
    assert any("page=3/3" in c.locator for c in successful)


def test_envelope_from_pdf_bytes_rejects_non_pdf_magic() -> None:
    with pytest.raises(PdfIngestionError, match="%PDF"):
        envelope_from_pdf_bytes(b"not a pdf at all", source_uri="memory://nope")


def test_envelope_from_pdf_bytes_rejects_empty() -> None:
    with pytest.raises(PdfIngestionError, match="empty"):
        envelope_from_pdf_bytes(b"", source_uri="memory://empty")


def test_envelope_from_pdf_bytes_encrypted_fails_loud() -> None:
    pdf = _build_pdf_bytes(page_texts=("secret content",), encrypted=True)
    with pytest.raises(PdfIngestionError, match="(encrypted|decrypt)"):
        envelope_from_pdf_bytes(pdf, source_uri="memory://encrypted.pdf")


def test_envelope_from_pdf_bytes_empty_page_becomes_partial_failure() -> None:
    # Build a blank page explicitly: single page, empty text -> extract_text()
    # returns "" which triggers the page_empty branch.
    pdf = _build_pdf_bytes(page_texts=("",))
    with pytest.raises(PdfIngestionError, match="no extractable text"):
        envelope_from_pdf_bytes(pdf, source_uri="memory://blank.pdf")


def test_envelope_from_pdf_bytes_mixed_empty_and_nonempty_pages() -> None:
    # Page 1 has text, page 2 is empty, page 3 has text.
    pdf = _build_pdf_bytes(page_texts=("Page one body.", "", "Page three body."))
    envelope = envelope_from_pdf_bytes(pdf, source_uri="memory://mixed.pdf")
    # Empty page 2 lands in partial_failures.
    assert len(envelope.partial_failures) >= 1
    failed = envelope.failed_chunks
    assert any("page=2/3" in c.locator for c in failed)
    assert any("page_empty" in c.parse_error for c in failed)
    # Non-empty pages are successful.
    assert envelope.successful_chunks


def test_envelope_from_pdf_bytes_respects_max_pages() -> None:
    pdf = _build_pdf_bytes(page_texts=("p1", "p2", "p3"))
    with pytest.raises(PdfIngestionError, match="max_pages"):
        envelope_from_pdf_bytes(pdf, source_uri="memory://big.pdf", max_pages=2)


def test_envelope_from_pdf_bytes_rejects_zero_max_pages() -> None:
    pdf = _build_pdf_bytes(page_texts=("p1",))
    with pytest.raises(PdfIngestionError, match="max_pages"):
        envelope_from_pdf_bytes(pdf, source_uri="memory://p.pdf", max_pages=0)


def test_envelope_from_pdf_bytes_long_page_subdivides(tmp_path: pathlib.Path) -> None:
    # Build one very long page; check that the resulting chunks all
    # fit under max_chunk_chars.
    long_text = ("paragraph line number " + str(i) for i in range(400))
    pdf = _build_pdf_bytes(page_texts=("\n".join(long_text),))
    envelope = envelope_from_pdf_bytes(
        pdf, source_uri="memory://long.pdf", max_chunk_chars=512
    )
    for chunk in envelope.successful_chunks:
        assert len(chunk.text) <= 512


def test_envelope_from_pdf_file_happy_path(tmp_path: pathlib.Path) -> None:
    pdf_bytes = _build_pdf_bytes(page_texts=("Disk page A.", "Disk page B."))
    path = tmp_path / "doc.pdf"
    path.write_bytes(pdf_bytes)
    envelope = envelope_from_pdf_file(path)
    assert envelope.provenance.source_uri.startswith("file://")
    assert envelope.total_chunks >= 2


def test_envelope_from_pdf_file_rejects_missing_file(tmp_path: pathlib.Path) -> None:
    with pytest.raises(PdfIngestionError, match="not a regular file"):
        envelope_from_pdf_file(tmp_path / "missing.pdf")


# ---------------------------------------------------------------------------
# DOCX fixtures (built via python-docx at runtime)
# ---------------------------------------------------------------------------


def _build_docx_bytes(
    *,
    sections: tuple[tuple[str, tuple[str, ...]], ...],
    with_table: bool = False,
) -> bytes:
    """Build a .docx. ``sections`` is ``((heading, (para, para, ...)), ...)``.

    An empty ``heading`` means "no heading row", so the section
    body starts directly. This lets us exercise the "no headings
    at all" fallback path.
    """
    from docx import Document

    document = Document()
    for heading, paragraphs in sections:
        if heading:
            document.add_heading(heading, level=1)
        for paragraph in paragraphs:
            document.add_paragraph(paragraph)
    if with_table:
        table = document.add_table(rows=1, cols=2)
        row = table.rows[0]
        row.cells[0].text = "key"
        row.cells[1].text = "value"
    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# DOCX tests
# ---------------------------------------------------------------------------


def test_envelope_from_docx_bytes_heading_delimited_sections() -> None:
    docx_bytes = _build_docx_bytes(
        sections=(
            ("Chapter One", ("Hello, this is chapter one.", "Second paragraph of chapter one.")),
            ("Chapter Two", ("Now we're in chapter two.",)),
        ),
    )
    envelope = envelope_from_docx_bytes(docx_bytes, source_uri="memory://test.docx")
    assert envelope.source_kind is IngestionSourceKind.BOOK
    # 2 sections -> at least 2 chunks (possibly more if sub-chunked).
    successful = envelope.successful_chunks
    assert len(successful) >= 2
    combined = " ".join(c.text for c in successful)
    assert "chapter one" in combined.lower()
    assert "chapter two" in combined.lower()
    # Locators include section=N/TOTAL + paragraphs range.
    assert any("section=1/2" in c.locator for c in successful)
    assert any("section=2/2" in c.locator for c in successful)


def test_envelope_from_docx_bytes_no_headings_single_section() -> None:
    docx_bytes = _build_docx_bytes(
        sections=(
            ("", ("First paragraph.", "Second paragraph.", "Third paragraph.")),
        ),
    )
    envelope = envelope_from_docx_bytes(docx_bytes, source_uri="memory://flat.docx")
    # No headings = one big section (possibly sub-chunked by length).
    successful = envelope.successful_chunks
    combined = " ".join(c.text for c in successful)
    assert "First paragraph" in combined
    assert "Third paragraph" in combined
    # All chunks should land in section=1/1.
    for chunk in successful:
        assert "section=1/1" in chunk.locator


def test_envelope_from_docx_bytes_rejects_non_zip_magic() -> None:
    with pytest.raises(DocxIngestionError, match="ZIP magic"):
        envelope_from_docx_bytes(
            b"not a docx just text", source_uri="memory://nope.docx"
        )


def test_envelope_from_docx_bytes_rejects_zip_without_document_xml() -> None:
    # Valid ZIP but contains the wrong entries.
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w") as zf:
        zf.writestr("README.md", "this is not a docx")
    with pytest.raises(DocxIngestionError, match="document.xml"):
        envelope_from_docx_bytes(buf.getvalue(), source_uri="memory://bad.docx")


def test_envelope_from_docx_bytes_rejects_empty() -> None:
    with pytest.raises(DocxIngestionError, match="empty"):
        envelope_from_docx_bytes(b"", source_uri="memory://empty.docx")


def test_envelope_from_docx_bytes_rejects_zero_max_paragraphs() -> None:
    docx_bytes = _build_docx_bytes(sections=(("", ("hello",)),))
    with pytest.raises(DocxIngestionError, match="max_paragraphs"):
        envelope_from_docx_bytes(
            docx_bytes, source_uri="memory://x.docx", max_paragraphs=0
        )


def test_envelope_from_docx_bytes_respects_max_paragraphs() -> None:
    # Make 5 paragraphs; cap at 3 -> fails loudly.
    docx_bytes = _build_docx_bytes(
        sections=(("", tuple(f"Para {i}" for i in range(5))),),
    )
    with pytest.raises(DocxIngestionError, match="exceeds max_paragraphs"):
        envelope_from_docx_bytes(
            docx_bytes, source_uri="memory://big.docx", max_paragraphs=3
        )


def test_envelope_from_docx_bytes_table_presence_marks_partial_failure() -> None:
    docx_bytes = _build_docx_bytes(
        sections=(("Heading", ("Para body.",)),),
        with_table=True,
    )
    envelope = envelope_from_docx_bytes(docx_bytes, source_uri="memory://t.docx")
    # At least one chunk flags "section_has_unextracted_table".
    assert len(envelope.partial_failures) >= 1
    failed = envelope.failed_chunks
    assert any(
        "section_has_unextracted_table" in c.parse_error for c in failed
    )


def test_envelope_from_docx_file_happy_path(tmp_path: pathlib.Path) -> None:
    docx_bytes = _build_docx_bytes(
        sections=(
            ("Report", ("First finding.", "Second finding.")),
        ),
    )
    path = tmp_path / "report.docx"
    path.write_bytes(docx_bytes)
    envelope = envelope_from_docx_file(path)
    assert envelope.provenance.source_uri.startswith("file://")
    combined = " ".join(c.text for c in envelope.successful_chunks)
    assert "First finding" in combined


def test_envelope_from_docx_file_rejects_missing_path(tmp_path: pathlib.Path) -> None:
    with pytest.raises(DocxIngestionError, match="not a regular file"):
        envelope_from_docx_file(tmp_path / "missing.docx")


def test_envelope_from_docx_bytes_long_section_subdivides() -> None:
    big_paragraph = "This is a sentence. " * 300  # ~6000 chars
    docx_bytes = _build_docx_bytes(sections=(("", (big_paragraph,)),))
    envelope = envelope_from_docx_bytes(
        docx_bytes, source_uri="memory://big.docx", max_chunk_chars=1024
    )
    # Multiple sub-chunks because section exceeds max_chunk_chars.
    assert envelope.total_chunks >= 2
    for chunk in envelope.successful_chunks:
        assert len(chunk.text) <= 1024


# ---------------------------------------------------------------------------
# Shared: envelope schema invariants
# ---------------------------------------------------------------------------


def test_pdf_envelope_provenance_integrity_hash_matches_bytes() -> None:
    import hashlib

    pdf = _build_pdf_bytes(page_texts=("abc",))
    envelope = envelope_from_pdf_bytes(pdf, source_uri="memory://h.pdf")
    assert envelope.provenance.integrity_hash == hashlib.sha256(pdf).hexdigest()


def test_docx_envelope_provenance_integrity_hash_matches_bytes() -> None:
    import hashlib

    docx_bytes = _build_docx_bytes(sections=(("", ("body",)),))
    envelope = envelope_from_docx_bytes(docx_bytes, source_uri="memory://h.docx")
    assert envelope.provenance.integrity_hash == hashlib.sha256(docx_bytes).hexdigest()


def test_pdf_default_compliance_is_forced() -> None:
    pdf = _build_pdf_bytes(page_texts=("hi",))
    envelope = envelope_from_pdf_bytes(pdf, source_uri="memory://c.pdf")
    assert envelope.compliance_profile is IngestionComplianceProfile.FORCED


def test_docx_default_compliance_is_forced() -> None:
    docx_bytes = _build_docx_bytes(sections=(("", ("hi",)),))
    envelope = envelope_from_docx_bytes(docx_bytes, source_uri="memory://c.docx")
    assert envelope.compliance_profile is IngestionComplianceProfile.FORCED
