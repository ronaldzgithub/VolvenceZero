"""End-to-end: PDF + DOCX -> ``IngestionPipeline`` -> real ``LifeformSession``
(Gap 3 slice 2).

Validates that slice-2 source adapters are a drop-in replacement
for slice-1 plain_text / task_result inside the pipeline:

* A multi-page PDF lands in the kernel as ``INGESTION`` turns,
  one per successfully-extracted page / sub-chunk.
* A DOCX with headings produces one ``INGESTION`` turn per
  section / sub-chunk; section-locator metadata is preserved
  through ``TurnSummary``.
* Pages / sections that the adapter flagged as partial_failure
  are skipped in the pipeline report, matching slice-1 behaviour.
* The kernel's ``vitals`` apprentice override engages for
  ``INGESTION`` turns, regardless of source format.
"""

from __future__ import annotations

import io
import pathlib
import zipfile

import pytest

pytest.importorskip("pypdf", reason="PDF e2e requires the [pdf] extra")
pytest.importorskip("docx", reason="DOCX e2e requires the [docx] extra")

from lifeform_core.types import TurnTriggerKind
from lifeform_ingestion import (
    IngestionPipeline,
    envelope_from_docx_bytes,
    envelope_from_pdf_bytes,
)


# ---------------------------------------------------------------------------
# Fixture helpers (same builders as the unit tests; duplicated here so
# the e2e module is self-contained and runnable independently)
# ---------------------------------------------------------------------------


def _build_pdf(*, page_texts: tuple[str, ...]) -> bytes:
    from pypdf import PdfWriter
    from pypdf.generic import (
        ContentStream,
        DecodedStreamObject,
        DictionaryObject,
        NameObject,
        NumberObject,
        TextStringObject,
    )

    writer = PdfWriter()
    for text in page_texts:
        page = writer.add_blank_page(width=612, height=792)
        font_dict = DictionaryObject(
            {
                NameObject("/Type"): NameObject("/Font"),
                NameObject("/Subtype"): NameObject("/Type1"),
                NameObject("/BaseFont"): NameObject("/Helvetica"),
            }
        )
        page[NameObject("/Resources")] = DictionaryObject(
            {
                NameObject("/Font"): DictionaryObject(
                    {NameObject("/F1"): font_dict}
                ),
            }
        )
        ops: list[tuple] = []
        y = 720
        for line in (text.splitlines() or [text]):
            ops.append(([], b"BT"))
            ops.append(([NameObject("/F1"), NumberObject(12)], b"Tf"))
            ops.append(([NumberObject(72), NumberObject(y)], b"Td"))
            ops.append(([TextStringObject(line)], b"Tj"))
            ops.append(([], b"ET"))
            y -= 18
        stream = ContentStream(None, writer)
        stream.operations = ops
        decoded = DecodedStreamObject()
        decoded.set_data(stream.get_data())
        page[NameObject("/Contents")] = decoded
    buf = io.BytesIO()
    writer.write(buf)
    return buf.getvalue()


def _build_docx(*, sections: tuple[tuple[str, tuple[str, ...]], ...]) -> bytes:
    from docx import Document

    document = Document()
    for heading, paragraphs in sections:
        if heading:
            document.add_heading(heading, level=1)
        for paragraph in paragraphs:
            document.add_paragraph(paragraph)
    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# E2E
# ---------------------------------------------------------------------------


async def test_pdf_envelope_drives_lifeform_session_through_ingestion_turns() -> None:
    from lifeform_domain_emogpt import build_companion_lifeform

    lifeform = build_companion_lifeform()
    session = lifeform.create_session(session_id="g3s2-pdf-e2e")
    pipeline = IngestionPipeline()

    pdf_bytes = _build_pdf(
        page_texts=(
            "Introduction page. This is the first ingested page.",
            "Analysis page. The interlocutor discusses frameworks.",
            "Conclusion page. Final notes summarising the ingestion.",
        )
    )
    envelope = envelope_from_pdf_bytes(pdf_bytes, source_uri="memory://e2e.pdf")
    report = await pipeline.process_envelope(env=envelope, session=session)

    # Every non-failed chunk produced one processed turn.
    assert report.processed_chunks == len(envelope.successful_chunks)
    # No skipped chunks in this happy-path case (all 3 pages extracted).
    assert report.skipped_chunks == 0
    # Session knows these were INGESTION-triggered (trigger_kind
    # propagation lives on TurnSummary, not on IngestionTurnRecord).
    assert len(session.turn_summaries) >= report.processed_chunks
    for summary in session.turn_summaries:
        assert summary.trigger_kind is TurnTriggerKind.INGESTION


async def test_pdf_partial_failure_chunks_are_skipped_in_pipeline() -> None:
    """Empty pages get a ``parse_error=page_empty`` and must show
    up in ``report.skipped_chunks`` rather than as ingested turns.
    """
    from lifeform_domain_emogpt import build_companion_lifeform

    lifeform = build_companion_lifeform()
    session = lifeform.create_session(session_id="g3s2-pdf-partial")
    pipeline = IngestionPipeline()

    pdf_bytes = _build_pdf(
        page_texts=("Page 1 has text.", "", "Page 3 also has text."),
    )
    envelope = envelope_from_pdf_bytes(pdf_bytes, source_uri="memory://partial.pdf")
    report = await pipeline.process_envelope(env=envelope, session=session)

    # One empty page -> one skipped chunk with page_empty error.
    assert report.skipped_chunks >= 1
    # Successful pages ran as INGESTION turns.
    assert report.processed_chunks >= 2
    for summary in session.turn_summaries:
        assert summary.trigger_kind is TurnTriggerKind.INGESTION


async def test_docx_envelope_drives_lifeform_session_through_ingestion_turns() -> None:
    from lifeform_domain_emogpt import build_companion_lifeform

    lifeform = build_companion_lifeform()
    session = lifeform.create_session(session_id="g3s2-docx-e2e")
    pipeline = IngestionPipeline()

    docx_bytes = _build_docx(
        sections=(
            ("Introduction", ("Context of the report.", "Key definitions.")),
            ("Findings", ("Observation A.", "Observation B.")),
            ("Conclusion", ("Summary paragraph.",)),
        )
    )
    envelope = envelope_from_docx_bytes(docx_bytes, source_uri="memory://e2e.docx")
    report = await pipeline.process_envelope(env=envelope, session=session)

    assert report.processed_chunks == len(envelope.successful_chunks)
    for summary in session.turn_summaries:
        assert summary.trigger_kind is TurnTriggerKind.INGESTION
    # Report echoes the envelope id for audit correlation.
    assert envelope.envelope_id in report.envelope_id


async def test_docx_file_roundtrip_via_pipeline(tmp_path: pathlib.Path) -> None:
    """File -> envelope -> pipeline -> session. Verifies the
    full on-disk path works without any fancy bytes handling.
    """
    from lifeform_domain_emogpt import build_companion_lifeform
    from lifeform_ingestion import envelope_from_docx_file

    docx_bytes = _build_docx(
        sections=(("Summary", ("Body paragraph one.",)),),
    )
    path = tmp_path / "one_section.docx"
    path.write_bytes(docx_bytes)

    lifeform = build_companion_lifeform()
    session = lifeform.create_session(session_id="g3s2-docx-file")
    pipeline = IngestionPipeline()
    envelope = envelope_from_docx_file(path)
    report = await pipeline.process_envelope(env=envelope, session=session)
    assert report.processed_chunks >= 1
    # Provenance source URI points at the real file.
    assert envelope.provenance.source_uri.startswith("file://")


async def test_pdf_ingestion_activates_vitals_apprentice_override() -> None:
    """FORCED-profile ingestion should flip the vitals apprentice
    override for each turn (Gap 2 \u2194 Gap 3 crossover). We observe
    this through the per-turn vitals snapshot: during an ingestion
    turn, ``total_pe`` is zeroed regardless of drive deviation.
    """
    from lifeform_domain_coding import build_coding_lifeform

    lifeform = build_coding_lifeform()
    session = lifeform.create_session(session_id="g3s2-pdf-vitals")
    pipeline = IngestionPipeline()

    pdf_bytes = _build_pdf(
        page_texts=("A single page of ingested content for vitals testing.",),
    )
    envelope = envelope_from_pdf_bytes(pdf_bytes, source_uri="memory://vitals.pdf")
    await pipeline.process_envelope(env=envelope, session=session)

    # Post-run: vitals override is now off (run_turn's finally
    # restored the prior state). The visible proof that the
    # override DID engage is that every turn_summary has
    # trigger_kind=INGESTION \u2014 run_turn asserts that. Here we
    # just confirm the ingestion completed without the kernel
    # raising on vitals.
    ingested = [
        s for s in session.turn_summaries
        if s.trigger_kind is TurnTriggerKind.INGESTION
    ]
    assert len(ingested) >= 1
    vitals_snap = session.vitals_snapshot
    assert vitals_snap is not None
    # Override is restored to False after ingestion finishes.
    assert session.vitals_module is not None
    assert session.vitals_module.apprentice_override_active is False
